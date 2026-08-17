"""
Yuklangan resume faylidan (PDF yoki DOCX) toza matnni ajratib olish.
"""
import io
from fastapi import HTTPException

import pdfplumber
from docx import Document


def extract_text_from_pdf(file_bytes: bytes) -> str:
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
    return "\n".join(text_parts).strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    document = Document(io.BytesIO(file_bytes))
    parts = [p.text for p in document.paragraphs if p.text.strip()]

    # Jadval ichidagi matnlarni ham olib qo'yamiz (ba'zi resumelar table asosida yoziladi)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())

    return "\n".join(parts).strip()


def extract_resume_text(filename: str, file_bytes: bytes) -> str:
    lower_name = filename.lower()

    if lower_name.endswith(".pdf"):
        text = extract_text_from_pdf(file_bytes)
    elif lower_name.endswith(".docx"):
        text = extract_text_from_docx(file_bytes)
    else:
        raise HTTPException(
            status_code=400,
            detail="Faqat .pdf yoki .docx formatidagi fayllar qo'llab-quvvatlanadi.",
        )

    if not text:
        raise HTTPException(
            status_code=422,
            detail="Fayldan matn topilmadi. Fayl skan qilingan rasm bo'lishi mumkin (matn qatlami yo'q).",
        )

    return text